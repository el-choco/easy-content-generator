from fastapi import FastAPI, HTTPException, Depends, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from typing import Optional
from datetime import timedelta, datetime
import google.generativeai as genai
import os
from io import BytesIO

from database import engine, get_db, Base
from models import User, Content, Template
from auth import (
    hash_password,
    verify_password,
    create_access_token,
    verify_token,
    get_current_user,
    ACCESS_TOKEN_EXPIRE_MINUTES
)

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch

Base.metadata.create_all(bind=engine)

app = FastAPI(title="Easy Content Generator", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["Content-Type", "Authorization", "Accept"],
    expose_headers=["*"],
    max_age=86400,
)

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

SUPPORTED_LANGUAGES = {
    "en": "English",
    "de": "German",
    "fr": "French",
    "es": "Spanish",
    "it": "Italian",
    "pt": "Portuguese",
    "nl": "Dutch",
    "ja": "Japanese",
    "zh": "Chinese",
    "ru": "Russian"
}

SUPPORTED_TONES = {
    "professional": "Professional - Formal, structured, business-appropriate tone",
    "casual": "Casual - Friendly, conversational, relaxed tone",
    "creative": "Creative - Imaginative, engaging, artistic tone",
    "technical": "Technical - Detailed, precise, specialized terminology"
}

DEFAULT_TEMPLATES = {
    "en": [
        {"name": "Blog Post", "category": "blog", "prompt": "Write a detailed and engaging blog post about {topic}. Include an introduction, main points, and a conclusion. Make it informative and SEO-friendly."},
        {"name": "Email Newsletter", "category": "email", "prompt": "Write a professional and engaging email newsletter about {topic}. Include a catchy subject line, introduction, main content, and a call to action."},
        {"name": "Instagram Post", "category": "social", "prompt": "Write a short and engaging Instagram post about {topic}. Include relevant hashtags and emojis. Keep it under 150 characters for the caption."},
        {"name": "Twitter Post", "category": "social", "prompt": "Write a concise and engaging Tweet about {topic}. Keep it under 280 characters and include relevant hashtags."},
        {"name": "Facebook Post", "category": "social", "prompt": "Write an engaging Facebook post about {topic}. Include a compelling headline and descriptive content that encourages shares and comments."},
        {"name": "TikTok Script", "category": "social", "prompt": "Write a short and entertaining TikTok script about {topic}. Make it trendy, engaging, and suitable for a 15-60 second video."},
        {"name": "Product Description", "category": "product", "prompt": "Write a compelling and detailed product description for {topic}. Highlight key features, benefits, and why customers should buy it."},
        {"name": "Job Description", "category": "hr", "prompt": "Write a professional and detailed job description for the position of {topic}. Include responsibilities, qualifications, and benefits."},
        {"name": "Press Release", "category": "pr", "prompt": "Write a professional press release about {topic}. Include a headline, introduction, main content, and closing statement."},
        {"name": "Newsletter Content", "category": "email", "prompt": "Write engaging newsletter content about {topic}. Make it informative and suitable for weekly distribution to subscribers."},
    ],
    "de": [
        {"name": "Blog Post", "category": "blog", "prompt": "Schreibe einen detaillierten und ansprechenden Blog-Post über {topic}. Include eine Einleitung, Hauptpunkte und ein Fazit. Mache ihn informativ und SEO-freundlich."},
        {"name": "E-Mail Newsletter", "category": "email", "prompt": "Schreibe einen professionellen und ansprechenden E-Mail Newsletter über {topic}. Included eine aussagekräftige Betreffzeile, Einleitung, Hauptinhalt und einen Call-to-Action."},
        {"name": "Instagram Post", "category": "social", "prompt": "Schreibe einen kurzen und ansprechenden Instagram-Post über {topic}. Included relevante Hashtags und Emojis. Halte ihn unter 150 Zeichen für die Caption."},
        {"name": "Twitter Post", "category": "social", "prompt": "Schreibe einen prägnanten und ansprechenden Tweet über {topic}. Halte ihn unter 280 Zeichen und included relevante Hashtags."},
        {"name": "Facebook Post", "category": "social", "prompt": "Schreibe einen ansprechenden Facebook-Post über {topic}. Included eine überzeugende Überschrift und beschreibenden Inhalt, der Shares und Kommentare fördert."},
        {"name": "TikTok Skript", "category": "social", "prompt": "Schreibe ein kurzes und unterhaltsames TikTok-Skript über {topic}. Mache es trendy, ansprechend und geeignet für ein 15-60 Sekunden Video."},
        {"name": "Produktbeschreibung", "category": "product", "prompt": "Schreibe eine überzeugende und detaillierte Produktbeschreibung für {topic}. Hebe Schlüsselfunktionen, Vorteile und Gründe zum Kauf hervor."},
        {"name": "Stellenbeschreibung", "category": "hr", "prompt": "Schreibe eine professionelle und detaillierte Stellenbeschreibung für die Position {topic}. Included Verantwortlichkeiten, Qualifikationen und Vorteile."},
        {"name": "Pressemitteilung", "category": "pr", "prompt": "Schreibe eine professionelle Pressemitteilung über {topic}. Included eine Überschrift, Einleitung, Hauptinhalt und eine Abschlusserklärung."},
        {"name": "Newsletter Inhalt", "category": "email", "prompt": "Schreibe ansprechenden Newsletter-Inhalt über {topic}. Mache ihn informativ und geeignet für die wöchentliche Verteilung an Abonnenten."},
    ],
    "fr": [
        {"name": "Article de Blog", "category": "blog", "prompt": "Écrivez un article de blog détaillé et attrayant sur {topic}. Incluez une introduction, les points principaux et une conclusion. Rendez-le informatif et convivial pour le SEO."},
        {"name": "Newsletter par E-mail", "category": "email", "prompt": "Écrivez une newsletter par e-mail professionnelle et attrayante sur {topic}. Incluez une ligne d'objet accrocheur, une introduction, le contenu principal et un appel à l'action."},
        {"name": "Post Instagram", "category": "social", "prompt": "Écrivez un post Instagram court et attrayant sur {topic}. Incluez des hashtags et des emojis pertinents. Gardez-le sous 150 caractères pour la légende."},
        {"name": "Tweet", "category": "social", "prompt": "Écrivez un tweet concis et attrayant sur {topic}. Gardez-le sous 280 caractères et incluez des hashtags pertinents."},
        {"name": "Post Facebook", "category": "social", "prompt": "Écrivez un post Facebook attrayant sur {topic}. Incluez un titre convaincant et un contenu descriptif qui encourage les partages et les commentaires."},
        {"name": "Script TikTok", "category": "social", "prompt": "Écrivez un script TikTok court et divertissant sur {topic}. Rendez-le tendance, attrayant et adapté à une vidéo de 15-60 secondes."},
        {"name": "Description de Produit", "category": "product", "prompt": "Écrivez une description de produit convaincante et détaillée pour {topic}. Mettez en évidence les caractéristiques clés, les avantages et pourquoi les clients devraient l'acheter."},
        {"name": "Offre d'Emploi", "category": "hr", "prompt": "Écrivez une offre d'emploi professionnelle et détaillée pour le poste de {topic}. Incluez les responsabilités, les qualifications et les avantages."},
        {"name": "Communiqué de Presse", "category": "pr", "prompt": "Écrivez un communiqué de presse professionnel sur {topic}. Incluez un titre, une introduction, le contenu principal et une déclaration de clôture."},
        {"name": "Contenu Newsletter", "category": "email", "prompt": "Écrivez un contenu de newsletter attrayant sur {topic}. Rendez-le informatif et adapté à la distribution hebdomadaire aux abonnés."},
    ],
    "es": [
        {"name": "Artículo de Blog", "category": "blog", "prompt": "Escriba un artículo de blog detallado y atractivo sobre {topic}. Incluya una introducción, puntos principales y una conclusión. Hágalo informativo y amigable para SEO."},
        {"name": "Boletín por Correo Electrónico", "category": "email", "prompt": "Escriba un boletín profesional y atractivo sobre {topic}. Incluya una línea de asunto atractiva, introducción, contenido principal y un llamado a la acción."},
        {"name": "Post de Instagram", "category": "social", "prompt": "Escriba una publicación corta y atractiva de Instagram sobre {topic}. Incluya hashtags y emojis relevantes. Manténgalo bajo 150 caracteres para el título."},
        {"name": "Tweet", "category": "social", "prompt": "Escriba un tweet conciso y atractivo sobre {topic}. Manténgalo bajo 280 caracteres e incluya hashtags relevantes."},
        {"name": "Post de Facebook", "category": "social", "prompt": "Escriba una publicación atractiva de Facebook sobre {topic}. Incluya un título convincente y contenido descriptivo que fomente comparticiones y comentarios."},
        {"name": "Script de TikTok", "category": "social", "prompt": "Escriba un script corto y entretenido de TikTok sobre {topic}. Hágalo tendencia, atractivo y adecuado para un video de 15-60 segundos."},
        {"name": "Descripción de Producto", "category": "product", "prompt": "Escriba una descripción de producto convincente y detallada para {topic}. Destaque las características clave, beneficios y por qué los clientes deberían comprarlo."},
        {"name": "Descripción de Puesto", "category": "hr", "prompt": "Escriba una descripción de puesto profesional y detallada para la posición de {topic}. Incluya responsabilidades, calificaciones y beneficios."},
        {"name": "Comunicado de Prensa", "category": "pr", "prompt": "Escriba un comunicado de prensa profesional sobre {topic}. Incluya un título, introducción, contenido principal y una declaración de cierre."},
        {"name": "Contenido de Boletín", "category": "email", "prompt": "Escriba contenido atractivo de boletín sobre {topic}. Hágalo informativo y adecuado para distribución semanal a suscriptores."},
    ],
    "it": [
        {"name": "Articolo di Blog", "category": "blog", "prompt": "Scrivi un articolo di blog dettagliato e accattivante su {topic}. Includi un'introduzione, i punti principali e una conclusione. Rendilo informativo e SEO-friendly."},
        {"name": "Newsletter Email", "category": "email", "prompt": "Scrivi una newsletter email professionale e accattivante su {topic}. Includi una riga di oggetto accattivante, introduzione, contenuto principale e una chiamata all'azione."},
        {"name": "Post Instagram", "category": "social", "prompt": "Scrivi un post Instagram breve e accattivante su {topic}. Includi hashtag ed emoji rilevanti. Mantienilo sotto 150 caratteri per la didascalia."},
        {"name": "Tweet", "category": "social", "prompt": "Scrivi un tweet conciso e accattivante su {topic}. Mantienilo sotto 280 caratteri e includi hashtag rilevanti."},
        {"name": "Post Facebook", "category": "social", "prompt": "Scrivi un post Facebook accattivante su {topic}. Includi un titolo convincente e contenuto descrittivo che incoraggi condivisioni e commenti."},
        {"name": "Script TikTok", "category": "social", "prompt": "Scrivi uno script TikTok breve e divertente su {topic}. Rendilo trendy, accattivante e adatto per un video di 15-60 secondi."},
        {"name": "Descrizione Prodotto", "category": "product", "prompt": "Scrivi una descrizione di prodotto convincente e dettagliata per {topic}. Evidenzia le caratteristiche chiave, i vantaggi e perché i clienti dovrebbero comprarlo."},
        {"name": "Descrizione Posizione", "category": "hr", "prompt": "Scrivi una descrizione di posizione professionale e dettagliata per il ruolo di {topic}. Includi responsabilità, qualifiche e vantaggi."},
        {"name": "Comunicato Stampa", "category": "pr", "prompt": "Scrivi un comunicato stampa professionale su {topic}. Includi un titolo, introduzione, contenuto principale e una dichiarazione di chiusura."},
        {"name": "Contenuto Newsletter", "category": "email", "prompt": "Scrivi un contenuto newsletter accattivante su {topic}. Rendilo informativo e adatto per la distribuzione settimanale ai sottoscritti."},
    ],
    "pt": [
        {"name": "Artigo de Blog", "category": "blog", "prompt": "Escreva um artigo de blog detalhado e envolvente sobre {topic}. Inclua uma introdução, pontos principais e uma conclusão. Torne-o informativo e amigável para SEO."},
        {"name": "Newsletter por Email", "category": "email", "prompt": "Escreva uma newsletter profissional e envolvente sobre {topic}. Inclua uma linha de assunto atraente, introdução, conteúdo principal e uma chamada para ação."},
        {"name": "Post Instagram", "category": "social", "prompt": "Escreva uma postagem curta e envolvente no Instagram sobre {topic}. Inclua hashtags e emojis relevantes. Mantenha-o abaixo de 150 caracteres para a legenda."},
        {"name": "Tweet", "category": "social", "prompt": "Escreva um tweet conciso e envolvente sobre {topic}. Mantenha-o abaixo de 280 caracteres e inclua hashtags relevantes."},
        {"name": "Post Facebook", "category": "social", "prompt": "Escreva uma postagem do Facebook envolvente sobre {topic}. Inclua um título convincente e conteúdo descritivo que encourage compartilhamentos e comentários."},
        {"name": "Script TikTok", "category": "social", "prompt": "Escreva um script TikTok breve e divertido sobre {topic}. Torne-o tendência, envolvente e adequado para um vídeo de 15-60 segundos."},
        {"name": "Descrição do Produto", "category": "product", "prompt": "Escreva uma descrição de produto convincente e detalhada para {topic}. Destaque os recursos principais, benefícios e por que os clientes devem comprá-lo."},
        {"name": "Descrição da Vaga", "category": "hr", "prompt": "Escreva uma descrição de vaga profissional e detalhada para o cargo de {topic}. Inclua responsabilidades, qualificações e benefícios."},
        {"name": "Comunicado à Imprensa", "category": "pr", "prompt": "Escreva um comunicado à imprensa profissional sobre {topic}. Inclua um título, introdução, conteúdo principal e uma declaração de encerramento."},
        {"name": "Conteúdo da Newsletter", "category": "email", "prompt": "Escreva um conteúdo atraente da newsletter sobre {topic}. Torne-o informativo e adequado para distribuição semanal aos assinantes."},
    ]
}

if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)
    try:
        models = genai.list_models()
        available_models = [m.name for m in models if 'generateContent' in m.supported_generation_methods]
        print(f"Available models: {available_models}")
        
        model_name = available_models[0] if available_models else 'models/gemini-2.5-flash'
        model = genai.GenerativeModel(model_name)
        print(f"Using model: {model_name}")
    except Exception as e:
        print(f"Error loading models: {e}")
        model = genai.GenerativeModel('models/gemini-2.5-flash')
else:
    model = None


def export_to_markdown(title: str, body: str) -> str:
    """Konvertiert Content zu Markdown"""
    return f"""# {title}

{body}

---
Generated with Easy Content Generator
"""

def export_to_docx(title: str, body: str) -> bytes:
    """Konvertiert Content zu Word (.docx)"""
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(body)
    doc.add_paragraph()
    doc.add_paragraph("Generated with Easy Content Generator")
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()

def export_to_pdf(title: str, body: str) -> bytes:
    """Konvertiert Content zu PDF"""
    output = BytesIO()
    doc = SimpleDocTemplate(output, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor='#1f2937',
        spaceAfter=30
    )
    story.append(Paragraph(title, title_style))
    story.append(Spacer(1, 0.3*inch))
    
    body_style = styles['BodyText']
    paragraphs = body.split('\n')
    for para in paragraphs:
        if para.strip():
            story.append(Paragraph(para, body_style))
        story.append(Spacer(1, 0.1*inch))
    
    story.append(Spacer(1, 0.5*inch))
    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=8,
        textColor='#9ca3af'
    )
    story.append(Paragraph("Generated with Easy Content Generator", footer_style))
    
    doc.build(story)
    output.seek(0)
    return output.getvalue()


# ============================================
# 🔐 AUTH ENDPOINTS
# ============================================

@app.post("/auth/register")
async def register(username: str, email: str, password: str, db: Session = Depends(get_db)):
    """Registriere einen neuen User"""
    
    if not username or not email or not password:
        raise HTTPException(status_code=400, detail="Username, email and password required")
    
    existing_user = db.query(User).filter(
        (User.username == username) | (User.email == email)
    ).first()
    
    if existing_user:
        raise HTTPException(status_code=400, detail="Username or email already exists")
    
    try:
        hashed_password = hash_password(password)
        new_user = User(
            username=username,
            email=email,
            hashed_password=hashed_password
        )
        db.add(new_user)
        db.commit()
        db.refresh(new_user)
        
        access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
        access_token = create_access_token(
            data={
                "sub": new_user.id,
                "is_admin": new_user.is_admin 
            },
            expires_delta=access_token_expires
        )        
        return {
            "access_token": access_token,
            "token_type": "bearer",
            "user": {
                "id": new_user.id,
                "username": new_user.username,
                "email": new_user.email,
                "is_admin": new_user.is_admin
            }
        }
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/auth/login")
async def login(username: str, password: str, db: Session = Depends(get_db)):
    """Login für einen User"""
    
    if not username or not password:
        raise HTTPException(status_code=400, detail="Username and password required")
    
    user = db.query(User).filter(User.username == username).first()
    
    if not user or not verify_password(password, user.hashed_password):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    if not user.is_active:
        raise HTTPException(status_code=403, detail="User is inactive")
    
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(
        data={
            "sub": user.id,
            "is_admin": user.is_admin 
        },
        expires_delta=access_token_expires
    )
    
    return {
        "access_token": access_token,
        "token_type": "bearer",
        "user": {
            "id": user.id,
            "username": user.username,
            "email": user.email,
            "is_admin": user.is_admin
        }
    }

@app.get("/auth/me")
async def get_me(current_user: User = Depends(get_current_user)):
    """Hole aktuellen User Info"""
    return {
        "id": current_user.id,
        "username": current_user.username,
        "email": current_user.email,
        "is_admin": current_user.is_admin,
        "is_active": current_user.is_active 
    }


# ============================================
# 📱 GENERAL ENDPOINTS
# ============================================

@app.get("/health")
async def health_check():
    return {"status": "ok", "timestamp": datetime.now().isoformat()}

@app.get("/languages")
async def get_languages():
    """Gibt alle unterstützten Sprachen zurück"""
    return {"languages": SUPPORTED_LANGUAGES}

@app.get("/tones")
async def get_tones():
    """Gibt alle unterstützten Tones zurück"""
    return {"tones": SUPPORTED_TONES}

@app.get("/")
async def root():
    return {
        "message": "Welcome to Easy Content Generator",
        "version": "1.0.0",
        "docs": "/docs"
    }


# ============================================
# 📝 GENERATION ENDPOINT
# ============================================

@app.post("/generate")
async def generate_content(
    prompt: str,
    language: str = "en",
    tone: str = "professional",
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Generiere Content (nur für authenticated users)"""
    
    if not model:
        raise HTTPException(status_code=500, detail="Gemini API not configured")
    
    if language not in SUPPORTED_LANGUAGES:
        raise HTTPException(status_code=400, detail=f"Unsupported language")
    
    if tone not in SUPPORTED_TONES:
        raise HTTPException(status_code=400, detail=f"Unsupported tone")
    
    try:
        language_name = SUPPORTED_LANGUAGES[language]
        tone_description = SUPPORTED_TONES[tone]
        
        enhanced_prompt = f"""Please answer in {language_name} with a {tone_description} tone.

{prompt}"""
        
        response = model.generate_content(enhanced_prompt)
        generated_text = response.text
        
        content = Content(
            title=prompt[:100],
            body=generated_text,
            language=language,
            tone=tone,
            status="published",  # ✅ Automatisch published
            owner_id=current_user.id
        )
        db.add(content)
        db.commit()
        db.refresh(content)
        
        return {
            "id": content.id,
            "prompt": prompt,
            "content": generated_text,
            "language": language,
            "tone": tone,
            "status": content.status,
            "created_at": content.created_at.isoformat()
        }
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ============================================
# 📚 CONTENT ENDPOINTS
# ============================================

@app.get("/content/{content_id}")
async def get_content(
    content_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Hole einen spezifischen Content"""
    content = db.query(Content).filter(
        Content.id == content_id,
        Content.owner_id == current_user.id
    ).first()
    
    if not content:
        raise HTTPException(status_code=404, detail="Content not found")
    
    return {
        "id": content.id,
        "title": content.title,
        "body": content.body,
        "language": content.language,
        "tone": content.tone,
        "status": content.status,
        "created_at": content.created_at.isoformat(),
        "updated_at": content.updated_at.isoformat() if content.updated_at else None
    }

@app.get("/history")
async def get_history(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Hole History des aktuellen Users (NUR published Content)"""
    contents = db.query(Content).filter(
        Content.owner_id == current_user.id,
        Content.status == "published"  # ✅ Nur published
    ).order_by(Content.created_at.desc()).all()
    
    return [
        {
            "id": content.id,
            "title": content.title,
            "body": content.body,
            "language": content.language,
            "tone": content.tone,
            "status": content.status,
            "created_at": content.created_at.isoformat()
        }
        for content in contents
    ]

@app.delete("/content/{content_id}")
async def delete_content(
    content_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Lösche einen Content"""
    content = db.query(Content).filter(
        Content.id == content_id,
        Content.owner_id == current_user.id
    ).first()
    
    if not content:
        raise HTTPException(status_code=404, detail="Content not found")
    
    db.delete(content)
    db.commit()
    
    return {"message": "Content deleted successfully"}

@app.put("/content/{content_id}")
async def update_content(
    content_id: int,
    title: str,
    body: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Update einen Content"""
    content = db.query(Content).filter(
        Content.id == content_id,
        Content.owner_id == current_user.id
    ).first()
    
    if not content:
        raise HTTPException(status_code=404, detail="Content not found")
    
    content.title = title
    content.body = body
    content.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(content)
    
    return {
        "id": content.id,
        "title": content.title,
        "body": content.body,
        "language": content.language,
        "tone": content.tone,
        "status": content.status
    }


# ============================================
# 🎯 DRAFT ENDPOINTS
# ============================================

@app.post("/drafts")
async def save_draft(
    title: str,
    body: str,
    language: str = "en",
    tone: str = "professional",
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Speichere einen Draft (unvollendeter Content)"""
    
    try:
        draft = Content(
            title=title if title.strip() else "Untitled Draft",
            body=body,
            language=language,
            tone=tone,
            status="draft",  # ✅ Status = draft
            owner_id=current_user.id
        )
        db.add(draft)
        db.commit()
        db.refresh(draft)
        
        return {
            "id": draft.id,
            "title": draft.title,
            "body": draft.body,
            "language": draft.language,
            "tone": draft.tone,
            "status": draft.status,
            "created_at": draft.created_at.isoformat(),
            "updated_at": draft.updated_at.isoformat()
        }
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/drafts")
async def get_drafts(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Hole alle Drafts des aktuellen Users"""
    
    drafts = db.query(Content).filter(
        Content.owner_id == current_user.id,
        Content.status == "draft"  # ✅ Nur Drafts
    ).order_by(Content.updated_at.desc()).all()
    
    return [
        {
            "id": draft.id,
            "title": draft.title,
            "body": draft.body,
            "language": draft.language,
            "tone": draft.tone,
            "status": draft.status,
            "created_at": draft.created_at.isoformat(),
            "updated_at": draft.updated_at.isoformat()
        }
        for draft in drafts
    ]


@app.put("/drafts/{draft_id}")
async def update_draft(
    draft_id: int,
    title: str = None,
    body: str = None,
    language: str = None,
    tone: str = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Update einen Draft"""
    
    draft = db.query(Content).filter(
        Content.id == draft_id,
        Content.owner_id == current_user.id,
        Content.status == "draft"  # ✅ Nur Drafts
    ).first()
    
    if not draft:
        raise HTTPException(status_code=404, detail="Draft not found")
    
    if title is not None:
        draft.title = title
    if body is not None:
        draft.body = body
    if language is not None:
        draft.language = language
    if tone is not None:
        draft.tone = tone
    
    draft.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(draft)
    
    return {
        "id": draft.id,
        "title": draft.title,
        "body": draft.body,
        "language": draft.language,
        "tone": draft.tone,
        "status": draft.status,
        "created_at": draft.created_at.isoformat(),
        "updated_at": draft.updated_at.isoformat()
    }


@app.put("/drafts/{draft_id}/publish")
async def publish_draft(
    draft_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Konvertiere Draft zu Published Content"""
    
    draft = db.query(Content).filter(
        Content.id == draft_id,
        Content.owner_id == current_user.id,
        Content.status == "draft"  # ✅ Nur Drafts
    ).first()
    
    if not draft:
        raise HTTPException(status_code=404, detail="Draft not found")
    
    draft.status = "published"  # ✅ Status ändern
    draft.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(draft)
    
    return {
        "id": draft.id,
        "title": draft.title,
        "body": draft.body,
        "language": draft.language,
        "tone": draft.tone,
        "status": draft.status,
        "created_at": draft.created_at.isoformat(),
        "updated_at": draft.updated_at.isoformat()
    }


@app.delete("/drafts/{draft_id}")
async def delete_draft(
    draft_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Lösche einen Draft"""
    
    draft = db.query(Content).filter(
        Content.id == draft_id,
        Content.owner_id == current_user.id,
        Content.status == "draft"  # ✅ Nur Drafts
    ).first()
    
    if not draft:
        raise HTTPException(status_code=404, detail="Draft not found")
    
    db.delete(draft)
    db.commit()
    
    return {"message": "Draft deleted successfully"}


# ============================================
# 📥 EXPORT ENDPOINTS
# ============================================

@app.get("/export/{content_id}/markdown")
async def export_markdown(
    content_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Exportiere Content als Markdown"""
    content = db.query(Content).filter(
        Content.id == content_id,
        Content.owner_id == current_user.id
    ).first()
    
    if not content:
        raise HTTPException(status_code=404, detail="Content not found")
    
    markdown_content = export_to_markdown(content.title, content.body)
    
    return FileResponse(
        BytesIO(markdown_content.encode()),
        media_type="text/markdown",
        filename=f"{content.title.replace(' ', '_')}.md"
    )

@app.get("/export/{content_id}/docx")
async def export_docx(
    content_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Exportiere Content als Word"""
    content = db.query(Content).filter(
        Content.id == content_id,
        Content.owner_id == current_user.id
    ).first()
    
    if not content:
        raise HTTPException(status_code=404, detail="Content not found")
    
    docx_bytes = export_to_docx(content.title, content.body)
    
    return FileResponse(
        BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{content.title.replace(' ', '_')}.docx"
    )

@app.get("/export/{content_id}/pdf")
async def export_pdf(
    content_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Exportiere Content als PDF"""
    content = db.query(Content).filter(
        Content.id == content_id,
        Content.owner_id == current_user.id
    ).first()
    
    if not content:
        raise HTTPException(status_code=404, detail="Content not found")
    
    pdf_bytes = export_to_pdf(content.title, content.body)
    
    return FileResponse(
        BytesIO(pdf_bytes),
        media_type="application/pdf",
        filename=f"{content.title.replace(' ', '_')}.pdf"
    )


# ============================================
# 📚 TEMPLATES ENDPOINTS
# ============================================

@app.get("/templates")
async def get_templates(
    language: str = "en",
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Hole alle Templates"""
    
    if language not in SUPPORTED_LANGUAGES:
        raise HTTPException(status_code=400, detail=f"Unsupported language")
    
    default_templates = DEFAULT_TEMPLATES.get(language, DEFAULT_TEMPLATES["en"])
    
    user_templates = db.query(Template).filter(
        Template.language == language,
        Template.is_default == False,
        Template.owner_id == current_user.id
    ).all()
    
    result = [
        {
            "id": f"default_{i}",
            "name": t["name"],
            "category": t["category"],
            "prompt": t["prompt"],
            "language": language,
            "is_default": True
        }
        for i, t in enumerate(default_templates)
    ]
    
    result.extend([
        {
            "id": t.id,
            "name": t.name,
            "category": t.category,
            "prompt": t.prompt,
            "language": t.language,
            "is_default": False,
            "owner_id": t.owner_id
        }
        for t in user_templates
    ])
    
    return result

@app.post("/templates")
async def create_template(
    name: str,
    category: str,
    prompt: str,
    language: str = "en",
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Erstelle ein neues Template"""
    
    if language not in SUPPORTED_LANGUAGES:
        raise HTTPException(status_code=400, detail=f"Unsupported language")
    
    template = Template(
        name=name,
        category=category,
        prompt=prompt,
        language=language,
        is_default=False,
        owner_id=current_user.id
    )
    db.add(template)
    db.commit()
    db.refresh(template)
    
    return {
        "id": template.id,
        "name": template.name,
        "category": template.category,
        "prompt": template.prompt,
        "language": template.language,
        "is_default": False
    }

@app.delete("/templates/{template_id}")
async def delete_template(
    template_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Lösche ein Template"""
    
    template = db.query(Template).filter(
        Template.id == template_id,
        Template.owner_id == current_user.id
    ).first()
    
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    
    if template.is_default:
        raise HTTPException(status_code=403, detail="Cannot delete default templates")
    
    db.delete(template)
    db.commit()
    
    return {"message": "Template deleted successfully"}

# ============================================
# 🔐 ADMIN ENDPOINTS
# ============================================

# Middleware für Admin-Check
def check_admin(current_user: User = Depends(get_current_user)):
    """Prüfe ob User Admin ist"""
    if not current_user.is_admin:
        raise HTTPException(status_code=403, detail="Admin access required")
    return current_user


# ============================================
# 📊 DASHBOARD
# ============================================

@app.get("/admin/dashboard")
async def admin_dashboard(
    admin_user: User = Depends(check_admin),
    db: Session = Depends(get_db)
):
    """Admin Dashboard mit Statistiken"""
    
    total_users = db.query(User).count()
    active_users = db.query(User).filter(User.is_active == True).count()
    admins = db.query(User).filter(User.is_admin == True).count()
    
    total_contents = db.query(Content).count()
    published_contents = db.query(Content).filter(Content.status == "published").count()
    draft_contents = db.query(Content).filter(Content.status == "draft").count()
    
    total_templates = db.query(Template).count()
    default_templates = db.query(Template).filter(Template.is_default == True).count()
    custom_templates = db.query(Template).filter(Template.is_default == False).count()
    
    # Top Languages
    from sqlalchemy import func
    top_languages = db.query(
        Content.language,
        func.count(Content.id).label("count")
    ).group_by(Content.language).order_by(func.count(Content.id).desc()).limit(5).all()
    
    # Top Tones
    top_tones = db.query(
        Content.tone,
        func.count(Content.id).label("count")
    ).group_by(Content.tone).order_by(func.count(Content.id).desc()).limit(5).all()
    
    return {
        "users": {
            "total": total_users,
            "active": active_users,
            "admins": admins
        },
        "content": {
            "total": total_contents,
            "published": published_contents,
            "drafts": draft_contents
        },
        "templates": {
            "total": total_templates,
            "default": default_templates,
            "custom": custom_templates
        },
        "top_languages": [
            {"language": lang, "count": count} for lang, count in top_languages
        ],
        "top_tones": [
            {"tone": tone, "count": count} for tone, count in top_tones
        ]
    }


# ============================================
# 👥 USER MANAGEMENT
# ============================================

@app.get("/admin/users")
async def get_all_users(
    admin_user: User = Depends(check_admin),
    db: Session = Depends(get_db)
):
    """Hole alle Users mit Statistiken"""
    
    users = db.query(User).all()
    
    result = []
    for user in users:
        content_count = db.query(Content).filter(Content.owner_id == user.id).count()
        draft_count = db.query(Content).filter(
            Content.owner_id == user.id,
            Content.status == "draft"
        ).count()
        published_count = db.query(Content).filter(
            Content.owner_id == user.id,
            Content.status == "published"
        ).count()
        
        result.append({
            "id": user.id,
            "username": user.username,
            "email": user.email,
            "is_active": user.is_active,
            "is_admin": user.is_admin,
            "created_at": user.created_at.isoformat(),
            "stats": {
                "total_content": content_count,
                "drafts": draft_count,
                "published": published_count
            }
        })
    
    return result


@app.get("/admin/users/{user_id}")
async def get_user_detail(
    user_id: int,
    admin_user: User = Depends(check_admin),
    db: Session = Depends(get_db)
):
    """Hole Detail-Info eines Users"""
    
    user = db.query(User).filter(User.id == user_id).first()
    
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    contents = db.query(Content).filter(Content.owner_id == user.id).all()
    templates = db.query(Template).filter(Template.owner_id == user.id).all()
    
    return {
        "id": user.id,
        "username": user.username,
        "email": user.email,
        "is_active": user.is_active,
        "is_admin": user.is_admin,
        "created_at": user.created_at.isoformat(),
        "contents": [
            {
                "id": c.id,
                "title": c.title,
                "status": c.status,
                "created_at": c.created_at.isoformat()
            }
            for c in contents
        ],
        "templates": [
            {
                "id": t.id,
                "name": t.name,
                "category": t.category
            }
            for t in templates
        ]
    }


@app.put("/admin/users/{user_id}/toggle-active")
async def toggle_user_active(
    user_id: int,
    admin_user: User = Depends(check_admin),
    db: Session = Depends(get_db)
):
    """Aktiviere/Deaktiviere User"""
    
    user = db.query(User).filter(User.id == user_id).first()
    
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    if user.id == admin_user.id:
        raise HTTPException(status_code=400, detail="Cannot toggle your own status")
    
    user.is_active = not user.is_active
    db.commit()
    db.refresh(user)
    
    return {
        "id": user.id,
        "username": user.username,
        "is_active": user.is_active
    }


@app.put("/admin/users/{user_id}/toggle-admin")
async def toggle_user_admin(
    user_id: int,
    admin_user: User = Depends(check_admin),
    db: Session = Depends(get_db)
):
    """Mache User zu Admin oder entferne Admin-Status"""
    
    user = db.query(User).filter(User.id == user_id).first()
    
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    if user.id == admin_user.id:
        raise HTTPException(status_code=400, detail="Cannot toggle your own admin status")
    
    user.is_admin = not user.is_admin
    db.commit()
    db.refresh(user)
    
    return {
        "id": user.id,
        "username": user.username,
        "is_admin": user.is_admin
    }


# ============================================
#           👥 USER MANAGEMENT 
# ============================================

@app.put("/admin/users/{user_id}")
async def edit_user(
    user_id: int,
    username: str = None,
    email: str = None,
    admin_user: User = Depends(check_admin),
    db: Session = Depends(get_db)
):
    """Editiere User (Username, Email)"""
    
    user = db.query(User).filter(User.id == user_id).first()
    
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    if username and username != user.username:
        existing = db.query(User).filter(User.username == username).first()
        if existing:
            raise HTTPException(status_code=400, detail="Username already exists")
        user.username = username
    
    if email and email != user.email:
        existing = db.query(User).filter(User.email == email).first()
        if existing:
            raise HTTPException(status_code=400, detail="Email already exists")
        user.email = email
    
    db.commit()
    db.refresh(user)
    
    return {
        "id": user.id,
        "username": user.username,
        "email": user.email,
        "message": "User updated successfully"
    }


@app.post("/admin/users/{user_id}/reset-password")
async def reset_user_password(
    user_id: int,
    new_password: str,
    admin_user: User = Depends(check_admin),
    db: Session = Depends(get_db)
):
    """Setze neues Password für User"""
    
    if len(new_password) < 6:
        raise HTTPException(status_code=400, detail="Password must be at least 6 characters")
    
    user = db.query(User).filter(User.id == user_id).first()
    
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    user.hashed_password = hash_password(new_password)
    db.commit()
    db.refresh(user)
    
    return {
        "id": user.id,
        "username": user.username,
        "message": "Password reset successfully"
    }


@app.post("/admin/users/bulk-delete")
async def bulk_delete_users(
    user_ids: list[int],
    admin_user: User = Depends(check_admin),
    db: Session = Depends(get_db)
):
    """Lösche mehrere User gleichzeitig"""
    
    if not user_ids:
        raise HTTPException(status_code=400, detail="No user IDs provided")
    
    if admin_user.id in user_ids:
        raise HTTPException(status_code=400, detail="Cannot delete yourself")
    
    users_to_delete = db.query(User).filter(User.id.in_(user_ids)).all()
    
    if not users_to_delete:
        raise HTTPException(status_code=404, detail="No users found")
    
    deleted_count = 0
    deleted_usernames = []
    
    for user in users_to_delete:
        deleted_usernames.append(user.username)
        db.delete(user)
        deleted_count += 1
    
    db.commit()
    
    return {
        "deleted_count": deleted_count,
        "deleted_users": deleted_usernames,
        "message": f"{deleted_count} users deleted successfully"
    }

# ============================================
# 📄 CONTENT MANAGEMENT
# ============================================

@app.get("/admin/contents")
async def get_all_contents(
    admin_user: User = Depends(check_admin),
    db: Session = Depends(get_db),
    status: str = None
):
    """Hole alle Contents (optional filterable nach Status)"""
    
    query = db.query(Content)
    
    if status:
        query = query.filter(Content.status == status)
    
    contents = query.order_by(Content.created_at.desc()).all()
    
    return [
        {
            "id": c.id,
            "title": c.title,
            "body": c.body[:100] + "..." if len(c.body) > 100 else c.body,
            "status": c.status,
            "language": c.language,
            "tone": c.tone,
            "owner_id": c.owner_id,
            "owner_username": db.query(User).filter(User.id == c.owner_id).first().username,
            "created_at": c.created_at.isoformat()
        }
        for c in contents
    ]


@app.get("/admin/contents/{content_id}")
async def get_content_detail(
    content_id: int,
    admin_user: User = Depends(check_admin),
    db: Session = Depends(get_db)
):
    """Hole vollständige Content-Info"""
    
    content = db.query(Content).filter(Content.id == content_id).first()
    
    if not content:
        raise HTTPException(status_code=404, detail="Content not found")
    
    owner = db.query(User).filter(User.id == content.owner_id).first()
    
    return {
        "id": content.id,
        "title": content.title,
        "body": content.body,
        "status": content.status,
        "language": content.language,
        "tone": content.tone,
        "owner": {
            "id": owner.id,
            "username": owner.username,
            "email": owner.email
        },
        "created_at": content.created_at.isoformat(),
        "updated_at": content.updated_at.isoformat() if content.updated_at else None
    }


@app.delete("/admin/contents/{content_id}")
async def admin_delete_content(
    content_id: int,
    admin_user: User = Depends(check_admin),
    db: Session = Depends(get_db)
):
    """Lösche Content als Admin"""
    
    content = db.query(Content).filter(Content.id == content_id).first()
    
    if not content:
        raise HTTPException(status_code=404, detail="Content not found")
    
    db.delete(content)
    db.commit()
    
    return {"message": "Content deleted"}

@app.post("/admin/contents/bulk-delete")
async def bulk_delete_contents(
    request: dict,
    admin_user: User = Depends(check_admin),
    db: Session = Depends(get_db)
):
    """Lösche mehrere Contents gleichzeitig"""
    
    content_ids = request.get("content_ids", [])
    
    if not content_ids:
        raise HTTPException(status_code=400, detail="No content IDs provided")
    
    contents_to_delete = db.query(Content).filter(Content.id.in_(content_ids)).all()
    
    if not contents_to_delete:
        raise HTTPException(status_code=404, detail="No contents found")
    
    deleted_count = 0
    deleted_titles = []
    
    for content in contents_to_delete:
        deleted_titles.append(content.title)
        db.delete(content)
        deleted_count += 1
    
    db.commit()
    
    return {
        "deleted_count": deleted_count,
        "deleted_contents": deleted_titles,
        "message": f"{deleted_count} contents deleted successfully"
    }


# ============================================
# 📚 TEMPLATES MANAGEMENT
# ============================================

@app.get("/admin/templates")
async def get_all_templates(
    admin_user: User = Depends(check_admin),
    db: Session = Depends(get_db)
):
    """Hole alle Templates (Default + Custom)"""
    
    templates = db.query(Template).order_by(Template.created_at.desc()).all()
    
    return [
        {
            "id": t.id,
            "name": t.name,
            "category": t.category,
            "language": t.language,
            "is_default": t.is_default,
            "owner_id": t.owner_id,
            "owner_username": db.query(User).filter(User.id == t.owner_id).first().username if t.owner_id else "System",
            "created_at": t.created_at.isoformat()
        }
        for t in templates
    ]


@app.delete("/admin/templates/{template_id}")
async def admin_delete_template(
    template_id: int,
    admin_user: User = Depends(check_admin),
    db: Session = Depends(get_db)
):
    """Lösche Custom Template als Admin"""
    
    template = db.query(Template).filter(Template.id == template_id).first()
    
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    
    if template.is_default:
        raise HTTPException(status_code=403, detail="Cannot delete default templates")
    
    db.delete(template)
    db.commit()
    
    return {"message": "Template deleted"}


# ============================================
# 🔧 SYSTEM MANAGEMENT
# ============================================

@app.get("/admin/system/health")
async def system_health(
    admin_user: User = Depends(check_admin),
    db: Session = Depends(get_db)
):
    """System Health Check"""
    
    try:
        db_test = db.query(User).count()
        db_status = "✅ OK"
    except Exception as e:
        db_status = f"❌ Error: {str(e)}"
    
    # Gemini API Status
    gemini_status = "✅ OK" if model else "❌ Not configured"
    
    return {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "database": db_status,
        "gemini_api": gemini_status,
        "version": "1.0.0"
    }


@app.get("/admin/system/stats")
async def system_stats(
    admin_user: User = Depends(check_admin),
    db: Session = Depends(get_db)
):
    """Detaillierte System-Statistiken"""
    
    from sqlalchemy import func
    
    return {
        "database": {
            "users": db.query(User).count(),
            "contents": db.query(Content).count(),
            "templates": db.query(Template).count()
        },
        "content_by_status": {
            "published": db.query(Content).filter(Content.status == "published").count(),
            "draft": db.query(Content).filter(Content.status == "draft").count()
        },
        "content_by_language": dict(
            db.query(Content.language, func.count(Content.id))
            .group_by(Content.language)
            .all()
        ),
        "content_by_tone": dict(
            db.query(Content.tone, func.count(Content.id))
            .group_by(Content.tone)
            .all()
        ),
        "templates_by_category": dict(
            db.query(Template.category, func.count(Template.id))
            .group_by(Template.category)
            .all()
        ),
        "timestamp": datetime.now().isoformat()
    }

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)